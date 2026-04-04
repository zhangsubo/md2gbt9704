#!/usr/bin/env python3
"""
GB/T 9704-2012 Markdown to DOCX Converter v3

Converts Markdown documents to DOCX following Chinese official document standards.
Usage: python md2docx.py input.md [output.docx]

字体跨平台适配：
- Windows: 方正小标宋简体, 仿宋_GB2312, 黑体, 楷体_GB2312
- Mac: 方正小标宋简体, 仿宋, STHeiti, 楷体
- Linux: 使用可用字体（如无则使用回退字体）
"""

import sys
import re
import platform
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Twips, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ==================== 字体配置 ====================
def get_available_fonts():
    """检测系统可用字体"""
    fonts = {
        'title': ['方正小标宋简体', '方正小标宋', 'SimHei', '黑体'],
        'body': ['仿宋_GB2312', '仿宋', 'FangSong', 'FangSong_GB2312'],
        'heading1': ['黑体', 'SimHei', 'Heiti'],
        'heading2': ['楷体_GB2312', '楷体', 'Kaiti', 'STKaiti'],
    }
    
    # 检查系统类型并调整
    system = platform.system()
    
    # 如果是 Linux，尝试使用 fc-list
    if system == 'Linux':
        try:
            import subprocess
            result = subprocess.run(['fc-list', ':lang=zh', '-f', '%{family}\n'],
                                    capture_output=True, text=True, timeout=5)
            available = set()
            for line in result.stdout.split('\n'):
                for families in line.split(','):
                    available.add(families.strip())
            
            for category in fonts:
                fonts[category] = [f for f in fonts[category] if f in available] or fonts[category]
        except Exception:
            pass
    
    return fonts

# 全局字体配置
FONTS = get_available_fonts()

def get_font(category):
    """获取指定类别的首选字体"""
    return FONTS.get(category, FONTS['body'])[0]


# ==================== GB/T 9704-2012 标准配置 ====================
PAGE_MARGIN = {
    "top": 3.7,    # cm
    "bottom": 3.5,  # cm
    "left": 2.7,    # cm (27mm)
    "right": 2.7,   # cm (27mm)
}

# 字号对照：二号=22pt，三号=16pt，小四号=12pt
LINE_SPACING = 28  # 固定值28磅
CHAR_INDENT = 32    # 2字符缩进 = 32pt (16pt * 2)

# 字体配置
FONT_SIZES = {
    "title": 22,       # 2号 = 22pt
    "heading1": 16,    # 3号 = 16pt
    "heading2": 16,    # 3号 = 16pt
    "heading3": 16,    # 3号 = 16pt
    "heading4": 16,    # 3号 = 16pt
    "body": 16,        # 3号 = 16pt
    "table": 12,       # 小四号 = 12pt
}


def set_page_setup(doc):
    """设置A4页面及边距"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4高度
    section.page_width = Cm(21.0)   # A4宽度
    section.top_margin = Cm(PAGE_MARGIN["top"])
    section.bottom_margin = Cm(PAGE_MARGIN["bottom"])
    section.left_margin = Cm(PAGE_MARGIN["left"])
    section.right_margin = Cm(PAGE_MARGIN["right"])


def set_font(run, font_cn, font_en, size, bold=False):
    """设置字体，同时清除原有格式（斜体、下划线、颜色等）"""
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    # 清除斜体
    run.font.italic = False
    # 清除下划线
    run.font.underline = False
    # 清除颜色（设置为黑色）
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 清除删除线
    run.font.strike = False
    run.font.double_strike = False
    # 清除上下标
    run.font.subscript = False
    run.font.superscript = False
    # 设置中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)


def apply_paragraph_style(para, font_name, font_size, bold=False,
                          align='left', indent=0, line_spacing=LINE_SPACING,
                          space_before=0, space_after=0):
    """应用统一的段落样式"""
    # 设置段落对齐
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    
    # 设置段落格式
    pf = para.paragraph_format
    # 重置段落缩进（重要：确保"文本之前缩进"为0）
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    
    # 首行缩进
    if indent > 0:
        pf.first_line_indent = Pt(indent)
    else:
        pf.first_line_indent = Pt(0)
    
    # 设置字体（使用新的 set_font 函数清除原有格式）
    for run in para.runs:
        set_font(run, font_name, font_name, font_size, bold)
    
    return para


def add_empty_paragraph(doc):
    """添加空行（用于标题与正文之间的空行）"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(LINE_SPACING)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    # 添加空 run 并设置字体为仿宋（跨平台适配）
    run = para.add_run('')
    set_font(run, get_font('body'), get_font('body'), 16, bold=False)
    return para


def clean_markdown_text(text):
    """清理 Markdown 标记，返回纯文本"""
    # 去除加粗标记 **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 去除斜体标记 *text*
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 去除无序列表标记 *、+、-
    text = re.sub(r'^[\s]*[\*\+\-][\s]+', '', text)
    # 去除有序列表标记（数字.）
    text = re.sub(r'^[\s]*\d+\.[\s]+', '', text)
    return text


def normalize_spacing(text):
    """
    规范化中英文之间的空格
    - 删除中文之间的空格
    - 中文与数字/字母之间加空格
    - 数字/字母与中文之间加空格
    - 删除中文标点后多余的空格
    """
    # 删除中文标点（、，。！？；：）后面的空格
    text = re.sub(r'([、，。！？；：])\s+', r'\1', text)
    # 删除纯中文或纯数字之间的空格
    # 中文和中文之间不留空格
    text = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', text)
    # 中文和数字/字母之间加空格
    text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z0-9])', r'\1 \2', text)
    text = re.sub(r'([a-zA-Z0-9])([\u4e00-\u9fff])', r'\1 \2', text)
    return text


def is_signature_line(text):
    """检测是否为落款行（单位名称或日期）"""
    # 日期格式：2026年1月3日 或 2026 年 1 月 3 日
    if re.match(r'^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日$', text.strip()):
        return 'date'
    # 单位名称：通常以"公司"、"局"、"委"、"处"、"部"、"办"、"中心"等结尾
    if re.match(r'^.{2,}(公司|局|委|处|部|办|中心|厅|院|所|站|组|会|署|集团|有限公司)$', text.strip()):
        return 'organization'
    return None


def parse_markdown(content):
    """
    解析Markdown内容，识别标题层级和表格。
    返回: [(level, text, ptype, extra), ...]
    level: 0=标题, 1-4=正文层次, None=正文
    ptype: title, heading1, heading2, heading3, heading4, body, table, signature
    extra: 表格数据（当ptype=table时）
    
    标题识别优先级：内容前缀 > Markdown标记
    - 一、 或 一级标题 → heading1
    - （一）或 二级标题 → heading2
    - 1. 或 三级标题 → heading3
    - （1）或 四级标题 → heading4
    """
    lines = content.split('\n')
    result = []
    i = 0
    
    while i < len(lines):
        stripped = lines[i].strip()
        
        if not stripped:
            i += 1
            continue
        
        # 检测表格开始（包含 | 的行，且下一行是分隔线）
        if '|' in stripped and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # 检查是否是表格分隔行（如 |---|---| 或 | --- | --- |）
            if re.match(r'^\|[\s\-:]+\|', next_line):
                # 开始解析表格
                table_rows = []
                # 解析表头
                header_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                table_rows.append(header_cells)
                i += 2  # 跳过分隔行
                
                # 解析数据行
                while i < len(lines):
                    row_line = lines[i].strip()
                    if '|' in row_line and row_line.startswith('|'):
                        row_cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        table_rows.append(row_cells)
                        i += 1
                    else:
                        break
                
                result.append((None, '', 'table', table_rows))
                continue
        
        # 主标题 (# 标题)
        if stripped.startswith('# ') and len(stripped) > 2:
            result.append((0, stripped[2:].strip(), 'title', None))
            i += 1
            continue
        
        # 去除 Markdown 标题标记，获取纯文本
        text = stripped
        for prefix in ['#### ', '### ', '## ', '# ']:
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
                break
        
        # 清理 Markdown 标记
        clean_text = clean_markdown_text(text)
        
        # 检测落款（单位名称或日期）
        sig_type = is_signature_line(clean_text)
        if sig_type:
            result.append((None, clean_text, 'signature', sig_type))
            i += 1
            continue
        
        # 根据内容前缀识别标题层级（优先级高于 Markdown 标记）
        if re.match(r'^[一二三四五六七八九十]+、', clean_text):
            result.append((1, clean_text, 'heading1', None))
        elif re.match(r'^（[一二三四五六七八九十]+）', clean_text) or re.match(r'^\([一二三四五六七八九十]+\)', clean_text):
            result.append((2, clean_text, 'heading2', None))
        elif re.match(r'^\d+\.', clean_text):
            result.append((3, clean_text, 'heading3', None))
        elif re.match(r'^（\d+）', clean_text) or re.match(r'^\(\d+\)', clean_text):
            result.append((4, clean_text, 'heading4', None))
        else:
            # 普通文本
            result.append((None, clean_text, 'body', None))
        
        i += 1
    
    return result


def create_table(doc, table_data):
    """
    创建 Word 表格
    table_data: [[cell1, cell2, ...], ...] 第一行为表头
    格式要求：
    - 表格线：直线，0.5磅
    - 表头：仿宋加粗，小四号（12pt）
    - 内容：仿宋不加粗，小四号（12pt）
    """
    if not table_data:
        return
    
    font_name = get_font('body')
    font_size = FONT_SIZES['table']  # 小四号 = 12pt
    
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 设置表格边框（0.5磅）
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5磅 = 4 (1/8 pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # 填充表格内容
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            
            # 清理单元格文本
            clean_text = clean_markdown_text(cell_text.strip())
            # 规范化空格
            clean_text = normalize_spacing(clean_text)
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(clean_text)
            
            # 表头加粗
            if row_idx == 0:
                set_font(run, font_name, font_name, font_size, bold=True)
            else:
                set_font(run, font_name, font_name, font_size, bold=False)
            
            # 设置段落格式
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(16)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
    
    # 表格后添加空行
    add_empty_paragraph(doc)
    
    return table


def create_signature_paragraph(doc, text, sig_type):
    """
    创建落款段落
    - 单位名称：右对齐
    - 日期：右对齐
    落款间距通过在正文中添加空行来控制
    """
    font_name = get_font('body')
    font_size = FONT_SIZES['body']
    
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, font_name, font_name, font_size, bold=False)
    
    # 设置右对齐
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 设置段落格式
    pf = para.paragraph_format
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_SPACING)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    return para


def create_styled_paragraph(doc, text, ptype):
    """根据段落类型创建并格式化段落"""
    # 规范化空格
    clean_text = normalize_spacing(text)
    
    # 标题前缀后不加空格（一、xxx、（一）xxx、1.xxx、（1）xxx）
    prefix_patterns = [
        r'^([一二三四五六七八九十]+、)(.*)',
        r'^(（[一二三四五六七八九十]+）)(.*)',
        r'^(\([一二三四五六七八九十]+\))(.*)',
        r'^(\d+\.)(.*)',
        r'^(（\d+）)(.*)',
        r'^(\(\d+\))(.*)',
    ]
    
    # 调试：打印原始文本
    # print(f"DEBUG: clean_text={clean_text!r}")
    for pattern in prefix_patterns:
        match = re.match(pattern, clean_text)
        if match:
            prefix, rest = match.groups()
            # 去除 rest 开头的空格
            rest = rest.lstrip()
            clean_text = prefix + rest
            break
    
    para = doc.add_paragraph()
    run = para.add_run(clean_text)
    
    if ptype == 'title':
        # 主标题：居中，方正小标宋简体，2号（22pt），不加粗
        font_name = get_font('title')
        apply_paragraph_style(para, font_name, FONT_SIZES["title"],
                             bold=False, align='center',
                             indent=0, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    elif ptype == 'heading1':
        # 一级标题：一、黑体，3号（16pt），加粗，首行缩进2字符
        font_name = get_font('heading1')
        apply_paragraph_style(para, font_name, FONT_SIZES["heading1"],
                             bold=True, align='left',
                             indent=CHAR_INDENT, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    elif ptype == 'heading2':
        # 二级标题：（一）楷体，3号（16pt），加粗，首行缩进2字符
        font_name = get_font('heading2')
        apply_paragraph_style(para, font_name, FONT_SIZES["heading2"],
                             bold=True, align='left',
                             indent=CHAR_INDENT, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    elif ptype == 'heading3':
        # 三级标题：1. 仿宋，3号（16pt），加粗，首行缩进2字符
        font_name = get_font('body')
        apply_paragraph_style(para, font_name, FONT_SIZES["heading3"],
                             bold=True, align='left',
                             indent=CHAR_INDENT, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    elif ptype == 'heading4':
        # 四级标题：（1）仿宋，3号（16pt），不加粗，首行缩进2字符
        font_name = get_font('body')
        apply_paragraph_style(para, font_name, FONT_SIZES["heading4"],
                             bold=False, align='left',
                             indent=CHAR_INDENT, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    else:  # body
        # 正文：仿宋，3号（16pt），两端对齐，首行缩进2字符
        font_name = get_font('body')
        apply_paragraph_style(para, font_name, FONT_SIZES["body"],
                             bold=False, align='justify',
                             indent=CHAR_INDENT, line_spacing=LINE_SPACING,
                             space_before=0, space_after=0)
    
    return para


def set_font(run, font_ascii, font_east_asia, size, bold=False):
    """设置 run 的字体"""
    run.font.name = font_ascii
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_east_asia)


def add_page_numbers(doc):
    """
    添加页码（GB/T 9704-2012 标准）
    - 四号宋体 14pt
    - 格式：— 页码 —（一字线）
    - 奇数页居右空一字，偶数页居左空一字
    """
    section = doc.sections[0]
    
    # 启用奇偶页页眉页脚（文档级）
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        settings_el = doc.settings._element
        if settings_el.find(qn('w:evenAndOddHeaders')) is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))
    
    section.odd_and_even_pages_header_footer = True
    section.footer_distance = Cm(0.7)
    
    odd_footer = section.footer
    even_footer = section.even_page_footer
    odd_footer.is_linked_to_previous = False
    even_footer.is_linked_to_previous = False
    
    for para in odd_footer.paragraphs:
        para.clear()
    for para in even_footer.paragraphs:
        para.clear()
    
    def _build_footer_line(footer, align, pad_fullwidth):
        """构建页脚行"""
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        
        para.alignment = align
        
        # 前导全角空格（空一字）
        if pad_fullwidth:
            run0 = para.add_run("　")
            set_font(run0, 'Times New Roman', 'Times New Roman', 14, bold=False)
        
        # 左一字线（带空格）
        run1 = para.add_run("— ")
        set_font(run1, 'Times New Roman', 'Times New Roman', 14, bold=False)
        
        # 页码域
        run2 = para.add_run()
        set_font(run2, 'Times New Roman', 'Times New Roman', 14, bold=False)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar1)
        
        run3 = para.add_run()
        set_font(run3, 'Times New Roman', 'Times New Roman', 14, bold=False)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run3._r.append(instrText)
        
        run4 = para.add_run()
        set_font(run4, 'Times New Roman', 'Times New Roman', 14, bold=False)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run4._r.append(fldChar2)
        
        run5 = para.add_run()
        set_font(run5, 'Times New Roman', 'Times New Roman', 14, bold=False)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar3)
        
        # 右一字线（带空格）
        run6 = para.add_run(" —")
        set_font(run6, 'Times New Roman', 'Times New Roman', 14, bold=False)
        
        # 末尾全角空格（空一字）
        if not pad_fullwidth:
            run7 = para.add_run("　")
            set_font(run7, 'Times New Roman', 'Times New Roman', 14, bold=False)
    
    # 奇数页居右空一字，偶数页居左空一字
    _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
    _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)


def convert_markdown_to_docx(input_path, output_path):
    """主转换函数"""
    # 读取输入文件
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建文档
    doc = Document()
    
    # 设置页面
    set_page_setup(doc)
    
    # 解析Markdown
    parsed = parse_markdown(content)
    
    # 处理段落
    prev_was_title = False
    prev_was_heading = False
    last_signature_type = None
    
    for item in parsed:
        if len(item) == 4:
            level, text, ptype, extra = item
        else:
            level, text, ptype = item
            extra = None
        
        # 处理表格
        if ptype == 'table' and extra:
            create_table(doc, extra)
            prev_was_title = False
            prev_was_heading = False
            continue
        
        # 处理落款
        if ptype == 'signature':
            # 单位名称前空两行
            if extra == 'organization':
                add_empty_paragraph(doc)
                add_empty_paragraph(doc)
            
            # 日期前空一行
            if extra == 'date':
                add_empty_paragraph(doc)
            
            create_signature_paragraph(doc, text, extra)
            last_signature_type = extra
            prev_was_title = False
            prev_was_heading = False
            continue
        
        if not text.strip():
            continue
        
        # 主标题之后添加空行
        if prev_was_title and ptype != 'title':
            add_empty_paragraph(doc)
        
        # 创建段落
        para = create_styled_paragraph(doc, text, ptype)
        
        # 更新状态
        prev_was_title = (ptype == 'title')
        prev_was_heading = ptype in ('heading1', 'heading2', 'heading3', 'heading4')
        last_signature_type = None
    
    # 添加页码
    add_page_numbers(doc)
    
    # 保存文档
    doc.save(output_path)
    print(f"✓ 已转换: {input_path}")
    print(f"  → {output_path}")
    print(f"  字体配置: {get_font('title')} / {get_font('body')}")


def main():
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <输入.md> [输出.docx]")
        print("示例: python md2docx.py input.md")
        print("      python md2docx.py input.md output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if not Path(input_file).exists():
        print(f"错误: 找不到输入文件: {input_file}")
        sys.exit(1)
    
    # 如果没有指定输出文件，自动生成 "原文件名_format.docx"
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        input_path = Path(input_file)
        output_file = str(input_path.parent / f"{input_path.stem}_format.docx")
    
    convert_markdown_to_docx(input_file, output_file)


if __name__ == '__main__':
    main()
