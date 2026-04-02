from docx import Document

doc = Document('/Users/zhangsubo/Code/md2gbt9704/test/数据资产成本筛选与分摊建议_format.docx')
section = doc.sections[0]

print('different_first_page_header_footer:', section.different_first_page_header_footer)
print()

print('默认页脚 (奇数页):')
if section.footer.paragraphs:
    print('  文本:', repr(section.footer.paragraphs[0].text))
    print('  runs数:', len(section.footer.paragraphs[0].runs))
else:
    print('  空')

print()
print('偶数页页脚:')
if section.even_page_footer and section.even_page_footer.paragraphs:
    print('  文本:', repr(section.even_page_footer.paragraphs[0].text))
    print('  runs数:', len(section.even_page_footer.paragraphs[0].runs))
else:
    print('  空')

print()
print('首页页脚:')
if section.first_page_footer and section.first_page_footer.paragraphs:
    print('  文本:', repr(section.first_page_footer.paragraphs[0].text))
else:
    print('  空')
